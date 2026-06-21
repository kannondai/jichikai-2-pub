"""
HTMLファイルの図をWord文書にエクスポートする
"""
import os
from pathlib import Path
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches

def html_to_word(html_path: str, output_path: str):
    """
    HTMLファイルをブラウザでレンダリングしてスクリーンショットを撮り、
    Word文書に挿入する
    
    Args:
        html_path: 入力HTMLファイルのパス
        output_path: 出力Wordファイルのパス
    """
    # 絶対パスに変換
    html_path = Path(html_path).resolve()
    output_path = Path(output_path).resolve()
    
    # 一時的なスクリーンショットパス
    screenshot_path = html_path.parent / f"{html_path.stem}_screenshot.png"
    
    print(f"HTMLファイルをレンダリング中: {html_path}")
    
    # PlaywrightでHTMLをレンダリングしてスクリーンショット
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page(viewport={'width': 1200, 'height': 1600})
        
        # ファイルURLとして開く
        file_url = f"file:///{str(html_path).replace(os.sep, '/')}"
        page.goto(file_url)
        
        # ページが完全に読み込まれるまで待機
        page.wait_for_load_state('networkidle')
        page.wait_for_timeout(500)  # レンダリング完了を確実にするため
        
        # SVG要素のスクリーンショットを撮る
        svg_element = page.locator('svg')
        if svg_element.count() > 0:
            svg_element.screenshot(path=str(screenshot_path))
        else:
            # SVGがない場合はページ全体のスクリーンショット
            page.screenshot(path=str(screenshot_path), full_page=True)
        
        browser.close()
    
    print(f"スクリーンショット保存: {screenshot_path}")
    
    # Word文書を作成
    doc = Document()
    
    # タイトルを追加
    doc.add_heading('第二自治会　会計のお金の流れ', 0)
    
    # 説明を追加
    doc.add_paragraph('収入 → 口座 → 支出 の流れ　（募金は口座を経由しません）')
    
    # 画像を追加（幅を6インチに設定）
    doc.add_picture(str(screenshot_path), width=Inches(7))
    
    # 説明セクションを追加
    doc.add_heading('従来からの変更点', level=1)
    doc.add_paragraph(
        '定期積立（会費3,600円中の600円）を臨時大口（百万円）の流れに合わせて、'
        '一般会計から特別会計に振り替える。',
        style='List Number'
    )
    
    # Word文書を保存
    doc.save(str(output_path))
    print(f"Word文書を保存しました: {output_path}")
    
    # 一時ファイルを削除
    if screenshot_path.exists():
        screenshot_path.unlink()
        print(f"一時ファイルを削除: {screenshot_path}")

if __name__ == '__main__':
    # 現在のディレクトリのaccount_flow.htmlを変換
    script_dir = Path(__file__).parent
    html_file = script_dir / 'account_flow.html'
    word_file = script_dir / 'account_flow.docx'
    
    if not html_file.exists():
        print(f"エラー: {html_file} が見つかりません")
        exit(1)
    
    html_to_word(str(html_file), str(word_file))
    print("\n完了!")
